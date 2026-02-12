"""Template management for modern DOCX output.

This module provides comprehensive DOCX template creation and management
with support for modern Word standards, configurable styling, and
robust error handling.

Features:
- Modern DOCX template generation with Office 2019+ compatibility
- Configurable page layouts, fonts, and styling
- Built-in Word styles for maximum compatibility
- XML-level customization for advanced features
- Comprehensive error handling and validation
"""

from __future__ import annotations

import contextlib
import logging
from pathlib import Path
from typing import Any, Literal, Optional, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches, Pt

from .config import DEFAULT_CONFIG, TemplateConfig
from .exceptions import TemplateError

# Configure logger
logger = logging.getLogger(__name__)

# Type definitions
PageSize = Literal["A4", "Letter"]


class DocxTemplateManager:
    """Manage creation and customization of modern DOCX templates.

    This class provides methods to create DOCX templates with modern styling,
    proper page layouts, and comprehensive style definitions that work well
    with Pandoc and other document processing tools.

    Example:
        Create a basic template:
        >>> manager = DocxTemplateManager()
        >>> template_path = manager.create_template("my_template.docx")

        Create with custom settings:
        >>> manager = DocxTemplateManager(
        ...     page_size="Letter",
        ...     body_font="Arial",
        ...     heading_font="Georgia"
        ... )
        >>> template_path = manager.create_template("custom_template.docx")
    """

    def __init__(
        self,
        *,
        config: Optional[TemplateConfig] = None,
        page_size: Optional[PageSize] = None,
        margin_cm: Optional[float] = None,
        body_font: Optional[str] = None,
        body_size_pt: Optional[int] = None,
        heading_font: Optional[str] = None,
        code_font: Optional[str] = None,
        code_size_pt: Optional[int] = None,
    ) -> None:
        """Initialize template manager with configuration.

        Args:
            config: Template configuration object. If None, uses default.
            page_size: Page size ("A4" or "Letter"). Overrides config if provided.
            margin_cm: Margin size in centimeters. Overrides config if provided.
            body_font: Font for body text. Overrides config if provided.
            body_size_pt: Body text size in points. Overrides config if provided.
            heading_font: Font for headings. Overrides config if provided.
            code_font: Font for code blocks. Overrides config if provided.
            code_size_pt: Code font size in points. Overrides config if provided.
        """
        # Use provided config or default
        base_config = config or DEFAULT_CONFIG.template

        # Override with any explicitly provided parameters
        self.page_size: str = (
            page_size if page_size is not None else base_config.page_size
        )
        self.margin_cm = margin_cm if margin_cm is not None else base_config.margin_cm
        self.body_font = body_font if body_font is not None else base_config.body_font
        self.body_size_pt = (
            body_size_pt if body_size_pt is not None else base_config.body_size_pt
        )
        self.heading_font = (
            heading_font if heading_font is not None else base_config.heading_font
        )
        self.code_font = code_font if code_font is not None else base_config.code_font
        self.code_size_pt = (
            code_size_pt if code_size_pt is not None else base_config.code_size_pt
        )

        # Validate page size
        if self.page_size not in ("A4", "Letter"):
            raise TemplateError(
                None, f"Invalid page size: {self.page_size}. Must be 'A4' or 'Letter'"
            )

    def create_template(
        self,
        output_path: Union[str, Path],
        *,
        add_sample_content: bool = False,
    ) -> Path:
        """Create a DOCX template file with configured settings.

        Args:
            output_path: Path where the template file will be created
            add_sample_content: Whether to include sample content for preview

        Returns:
            Path to the created template file

        Raises:
            TemplateError: If template creation fails
        """
        return self._create_template_file(
            output_path, add_sample_content=add_sample_content
        )

    @classmethod
    def create_modern_template(
        cls,
        output_path: Union[str, Path],
        *,
        add_sample: bool = False,
        **kwargs: Any,
    ) -> Path:
        """Create a modern DOCX template using default or custom settings.

        This convenience class method creates a template manager instance
        and generates a template file in one step.

        Args:
            output_path: Path where the template will be created
            add_sample: Whether to include sample content
            **kwargs: Additional configuration parameters (see __init__)

        Returns:
            Path to the created template file

        Raises:
            TemplateError: If template creation fails

        Example:
            >>> template_path = DocxTemplateManager.create_modern_template(
            ...     "modern.docx",
            ...     page_size="Letter",
            ...     body_font="Arial"
            ... )
        """
        try:
            manager = cls(**kwargs)
            return manager._create_template_file(
                output_path, add_sample_content=add_sample
            )
        except Exception as e:
            raise TemplateError(
                str(output_path), f"Failed to create template: {e}"
            ) from e

    def _create_template_file(
        self, output_path: Union[str, Path], *, add_sample_content: bool = False
    ) -> Path:
        """Internal method to create the template file.

        Args:
            output_path: Path where template will be created
            add_sample_content: Whether to add sample content

        Returns:
            Path to created template

        Raises:
            TemplateError: If creation fails
        """
        try:
            output_path = Path(output_path)

            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)

            logger.info("Creating DOCX template: %s", output_path)

            # Create new document
            doc = Document()

            # Configure document structure and styles
            self._configure_page_layout(doc)
            self._configure_core_styles(doc)

            # Add sample content if requested
            if add_sample_content:
                self._add_sample_content(doc)

            # Set modern compatibility mode for better Office support
            self._set_compatibility_mode_xml(doc, mode="16")

            # Save the document
            doc.save(str(output_path))

            logger.info("Successfully created template: %s", output_path)
            return output_path

        except Exception as e:
            raise TemplateError(
                str(output_path), f"Failed to create template file: {e}"
            ) from e

    # ---------- Layout & styles ----------

    def _configure_page_layout(self, doc: Any) -> None:
        """Configure page size and margins."""
        for section in doc.sections:
            # Margins
            section.top_margin = Cm(self.margin_cm)
            section.bottom_margin = Cm(self.margin_cm)
            section.left_margin = Cm(self.margin_cm)
            section.right_margin = Cm(self.margin_cm)

            # Page size
            if self.page_size == "A4":
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
            else:  # Letter
                section.page_width = Inches(8.5)
                section.page_height = Inches(11.0)

    def _configure_core_styles(self, doc: Any) -> None:
        """Configure Normal, Heading 1..6, and a Code Block paragraph style."""
        styles = doc.styles

        # Normal
        normal = styles["Normal"]
        normal.font.name = self.body_font
        normal.font.size = Pt(self.body_size_pt)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        # Headings (use built-in names so TOC and outline work everywhere)
        # Sizes chosen for readability; adjust as desired
        heading_specs = [
            ("Heading 1", 18, True, 12, 6),
            ("Heading 2", 14, True, 10, 4),
            ("Heading 3", 12, True, 8, 3),
            ("Heading 4", 11, True, 6, 3),
            ("Heading 5", 11, False, 6, 3),
            ("Heading 6", 11, False, 6, 3),
        ]
        for name, size_pt, bold, space_before_pt, space_after_pt in heading_specs:
            style = (
                styles[name]
                if name in styles
                else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            )
            style.font.name = self.heading_font
            style.font.size = Pt(size_pt)
            style.font.bold = bold
            p = style.paragraph_format
            p.space_before = Pt(space_before_pt)
            p.space_after = Pt(space_after_pt)
            p.keep_with_next = True  # keep heading with the following paragraph

        # Code Block paragraph style (not built-in; keep the name stable)
        code_name = "Code Block"
        if code_name in styles:
            code_style = styles[code_name]
        else:
            code_style = styles.add_style(code_name, WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = self.code_font
        code_style.font.size = Pt(self.code_size_pt)
        cp = code_style.paragraph_format
        cp.left_indent = Cm(0.75)
        cp.space_before = Pt(6)
        cp.space_after = Pt(6)

    def _add_sample_content(self, doc: Any) -> None:
        """Insert minimal content for preview/testing the template styles."""
        doc.add_heading("Heading 1", level=1)
        doc.add_paragraph(
            "Body text under Heading 1. Replace or remove this sample content."
        )
        doc.add_heading("Heading 2", level=2)
        doc.add_paragraph("Body text under Heading 2.")
        doc.add_heading("Heading 3", level=3)
        doc.add_paragraph("Body text under Heading 3.")
        doc.add_paragraph("Sample code paragraph:", style="Normal")
        doc.add_paragraph("for i in range(3): print(i)", style="Code Block")

    # ---------- Low-level XML helpers (best-effort) ----------

    def _set_compatibility_mode_xml(self, doc: Any, *, mode: str = "16") -> None:
        """Set Word compatibility mode via settings.xml (best-effort).

        Notes:
        - python-docx does not expose this directly; we edit the XML tree.
        - Common values: 14 (Word 2010), 15 (Word 2013), 16 (Word 2016/2019+).
        - This hints Word rendering behavior for better modern compatibility.
        """
        with contextlib.suppress(Exception):
            # Find settings part
            settings_part = next(
                (
                    part
                    for part in doc.part._child_parts
                    if "settings" in str(part.partname)
                ),
                None,
            )

            if not settings_part:
                return

            settings = settings_part._element
            nsmap = settings.nsmap or {}
            w_ns = nsmap.get(
                "w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            )
            w = f"{{{w_ns}}}"

            # Find or create compat element
            compat = settings.find(f"{w}compat")
            if compat is None:
                compat = settings.makeelement(f"{w}compat")
                settings.append(compat)

            # Update or create compatibilityMode setting
            existing = compat.findall(f"{w}compatSetting")
            target = next(
                (el for el in existing if el.get(f"{w}name") == "compatibilityMode"),
                None,
            )

            if target is None:
                target = settings.makeelement(f"{w}compatSetting")
                compat.append(target)

            target.set(f"{w}name", "compatibilityMode")
            target.set(f"{w}uri", "http://schemas.microsoft.com/office/word")
            target.set(f"{w}val", mode)

    # ---------- Static method for backward compatibility ----------

    @classmethod
    def create_default_template(cls, output_path: str | Path) -> Path:
        """Create a modern DOCX template with default settings (backward compatibility)."""
        return cls.create_modern_template(output_path, add_sample=True)
