"""Integration tests for end-to-end functionality.

端到端集成测试，确保整个转换流程正常工作。
"""

import pytest
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from markdown2docx import MarkdownToDocxConverter, DocxTemplateManager


@pytest.fixture
def complex_markdown():
    """Complex markdown content for integration testing."""
    return """# Integration Test Document

This document tests the complete conversion pipeline.

## Text Formatting

**Bold text**, *italic text*, ~~strikethrough text~~, and `inline code`.

Superscript: E=mc^2^
Subscript: H~2~O

## Lists

### Unordered List
- First item
- Second item
  - Nested item
  - Another nested item
- Third item

### Ordered List
1. First step
2. Second step
3. Third step

### Task List
- [x] Completed task
- [ ] Pending task
- [ ] Another pending task

## Code Blocks

```python
def integration_test():
    \"\"\"Test function for integration testing.\"\"\"
    print("Integration test successful!")
    return True

# Call the function
result = integration_test()
```

```javascript
// JavaScript example
function testIntegration() {
    console.log("JavaScript integration test");
    return { status: "success" };
}
```

## Tables

| Feature | Status | Notes |
|---------|:------:|-------|
| Headers | ✅ | Working correctly |
| Lists | ✅ | All types supported |
| Code | ✅ | Syntax highlighting |
| Tables | ✅ | Complex layouts |

## Block Quotes

> This is a block quote with important information.
> 
> It can span multiple paragraphs and contain **formatting**.

## Links and References

Visit [GitHub](https://github.com) for more information.

This document has footnotes[^1] for additional context.

[^1]: This is a footnote with detailed information.

## Mathematical Expressions

The quadratic formula: $ax^2 + bx + c = 0$

## Conclusion

This integration test covers all major Markdown features to ensure comprehensive conversion capability.
"""


def test_full_conversion_pipeline(complex_markdown):
    """Test complete conversion pipeline from markdown to DOCX."""
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "integration_test.md"
        output_path = tmpdir_path / "integration_output.docx"
        
        input_path.write_text(complex_markdown)
        
        converter = MarkdownToDocxConverter()
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify document structure
        doc = Document(output_path)
        
        # Check that we have content
        assert len(doc.paragraphs) > 10
        
        # Check for headings
        headings = [p for p in doc.paragraphs if p.style and 'heading' in p.style.name.lower()]
        assert len(headings) >= 5  # Should have multiple heading levels
        
        # Check for tables
        assert len(doc.tables) >= 1


def test_template_based_conversion(complex_markdown):
    """Test conversion using custom template."""
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        template_path = tmpdir_path / "custom_template.docx"
        input_path = tmpdir_path / "test.md"
        output_path = tmpdir_path / "output.docx"
        
        # Create template
        DocxTemplateManager.create_modern_template(template_path)
        
        # Create input file
        input_path.write_text(complex_markdown)
        
        # Convert with template
        converter = MarkdownToDocxConverter(reference_doc=template_path)
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify template was used
        doc = Document(output_path)
        
        # Check margins (should match template)
        section = doc.sections[0]
        assert section.top_margin.emu == 914400  # 1 inch


def test_toc_generation(complex_markdown):
    """Test table of contents generation."""
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "test.md"
        output_path = tmpdir_path / "output.docx"
        
        input_path.write_text(complex_markdown)
        
        converter = MarkdownToDocxConverter()
        result = converter.convert(
            input_path, 
            output_path,
            **{'--toc': True, '--toc-depth': 2}
        )
        
        assert result == output_path
        assert output_path.exists()


def test_multilingual_conversion():
    """Test conversion of multilingual content."""
    multilingual_content = """# Multilingual Test 多语言测试

This document contains multiple languages.

## English Section

This is English content with **bold** and *italic* text.

## Chinese Section 中文部分

这是中文内容，包含**粗体**和*斜体*文本。

## Mixed Content 混合内容

This paragraph contains both English and 中文 characters in the same line.

### Table with Multiple Languages

| Language | Example | Status |
|----------|---------|--------|
| English | Hello World | ✅ |
| 中文 | 你好世界 | ✅ |
| Español | Hola Mundo | ✅ |
"""
    
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "multilingual.md"
        output_path = tmpdir_path / "multilingual_output.docx"
        
        input_path.write_text(multilingual_content)
        
        converter = MarkdownToDocxConverter()
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify multilingual content is preserved
        doc = Document(output_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert '多语言测试' in full_text
        assert '中文部分' in full_text
        assert '你好世界' in full_text
        assert 'Hola Mundo' in full_text


def test_error_recovery():
    """Test error handling and recovery."""
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        
        # Test with nonexistent input file
        converter = MarkdownToDocxConverter()
        
        with pytest.raises(FileNotFoundError):
            converter.convert("nonexistent.md")
        
        # Test with invalid template
        invalid_template = tmpdir_path / "invalid.docx"
        converter_with_template = MarkdownToDocxConverter(reference_doc=invalid_template)
        
        # Should still work without template
        input_path = tmpdir_path / "test.md"
        input_path.write_text("# Test")
        
        result = converter_with_template.convert(input_path)
        assert result.exists()


def test_large_document_conversion():
    """Test conversion of large documents."""
    # Generate large markdown content
    large_content = "# Large Document Test\n\n"
    
    for i in range(100):
        large_content += f"""## Section {i+1}

This is section {i+1} with some content. It contains **bold text**, *italic text*, and `inline code`.

### Subsection {i+1}.1

More content here with a list:

- Item 1
- Item 2
- Item 3

```python
def section_{i+1}_function():
    return "Section {i+1} code"
```

"""
    
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "large_document.md"
        output_path = tmpdir_path / "large_output.docx"
        
        input_path.write_text(large_content)
        
        converter = MarkdownToDocxConverter()
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify document has expected structure
        doc = Document(output_path)
        assert len(doc.paragraphs) > 500  # Should have many paragraphs


def test_special_characters_conversion():
    """Test conversion of documents with special characters."""
    special_content = """# Special Characters Test

## Currency Symbols
¥ $ € £ ₽ ₹ ₩ ₪

## Mathematical Symbols
∑ ∏ ∫ ∂ ∇ ∞ ± × ÷ ≠ ≤ ≥

## Arrow Symbols
← → ↑ ↓ ↔ ↕ ⇐ ⇒ ⇔

## Quotation Marks
"English quotes" 'Single quotes' « French quotes » „German quotes"

## Emoji and Unicode
🚀 🎉 ✅ ❌ 📝 💡 🔧 ⚡

## Special Punctuation
§ ¶ † ‡ • ‰ ′ ″ ‴ ※
"""
    
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "special_chars.md"
        output_path = tmpdir_path / "special_output.docx"
        
        input_path.write_text(special_content)
        
        converter = MarkdownToDocxConverter()
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify special characters are preserved
        doc = Document(output_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert '¥' in full_text
        assert '∑' in full_text
        assert '←' in full_text
        assert '🚀' in full_text