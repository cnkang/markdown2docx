"""Tests for DOCX style handling.

DOCX样式处理测试，确保Markdown标题正确映射到Word标题样式。
"""

import pytest
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from markdown2docx.converter import MarkdownToDocxConverter
from markdown2docx.templates import DocxTemplateManager


@pytest.fixture
def heading_markdown():
    """Markdown with various heading levels (各级标题的Markdown示例)."""
    return """# Heading 1 标题1

This is body content. 这是正文内容。

## Heading 2 标题2

This is more body content. 这是更多正文内容。

### Heading 3 标题3

Content under third level heading. 这是第三级标题下的内容。

#### Heading 4 标题4

Content under fourth level heading. 这是第四级标题下的内容。

##### Heading 5 标题5

Content under fifth level heading. 这是第五级标题下的内容。

###### Heading 6 标题6

Content under sixth level heading. 这是第六级标题下的内容。
"""


def test_heading_styles_mapping(heading_markdown):
    """Test that markdown headings map to correct Word heading styles.
    
    测试Markdown标题是否正确映射到Word标题样式。
    """
    with TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "test_headings.md"
        template_path = Path(tmpdir) / "template.docx"
        output_path = Path(tmpdir) / "output.docx"
        
        input_path.write_text(heading_markdown)
        DocxTemplateManager.create_modern_template(template_path)
        
        converter = MarkdownToDocxConverter(reference_doc=template_path)
        converter.convert(input_path, output_path)
        
        # Analyze the generated document (分析生成的文档)
        doc = Document(output_path)
        heading_styles = []
        
        for paragraph in doc.paragraphs:
            if paragraph.style and 'heading' in paragraph.style.name.lower():
                heading_styles.append((paragraph.text, paragraph.style.name))
        
        # Verify correct heading mappings (验证正确的标题映射)
        expected_patterns = [
            ('Heading 1', 'Heading 1'),
            ('Heading 2', 'Heading 2'),
            ('Heading 3', 'Heading 3'),
            ('Heading 4', 'Heading 4'),
            ('Heading 5', 'Heading 5'),
            ('Heading 6', 'Heading 6')
        ]
        
        assert len(heading_styles) == 6
        for i, (expected_pattern, expected_style) in enumerate(expected_patterns):
            actual_text, actual_style = heading_styles[i]
            assert expected_pattern in actual_text  # Text should contain the heading pattern
            assert actual_style == expected_style


def test_template_has_all_heading_styles():
    """Test that created template contains all heading styles.
    
    测试创建的模板包含所有标题样式。
    """
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path)
        
        doc = Document(template_path)
        style_names = [style.name for style in doc.styles]
        
        # Check that all heading styles exist (检查所有标题样式是否存在)
        for i in range(1, 7):
            assert f'Heading {i}' in style_names


def test_pandoc_styles_extension():
    """Test that converter uses +styles extension.
    
    测试转换器使用+styles扩展。
    """
    converter = MarkdownToDocxConverter()
    args = converter._get_modern_docx_args()
    
    # Check that docx+styles is used (检查是否使用docx+styles)
    assert '-t' in args
    docx_format_index = args.index('-t') + 1
    assert 'docx+styles' in args[docx_format_index]


def diagnose_docx_styles(docx_path):
    """Helper function to diagnose DOCX styles.
    
    诊断DOCX样式的辅助函数，用于分析文档中的样式使用情况。
    """
    doc = Document(docx_path)
    
    styles_info = {
        'available_styles': [style.name for style in doc.styles],
        'paragraph_styles': [],
        'heading_count': 0
    }
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            style_name = paragraph.style.name if paragraph.style else "No Style"
            styles_info['paragraph_styles'].append({
                'index': i + 1,
                'text': paragraph.text[:50] + '...' if len(paragraph.text) > 50 else paragraph.text,
                'style': style_name
            })
            
            if paragraph.style and 'heading' in paragraph.style.name.lower():
                styles_info['heading_count'] += 1
    
    return styles_info