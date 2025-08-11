"""Tests for template system functionality.

模板系统测试，确保DOCX模板创建和使用功能正常。
"""

import pytest
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import sys
from pathlib import Path
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from markdown2docx.templates import DocxTemplateManager
from markdown2docx.converter import MarkdownToDocxConverter


def test_create_modern_template():
    """Test modern template creation with all required styles."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "modern_template.docx"
        
        result = DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        
        assert result == template_path
        assert template_path.exists()
        
        # Verify template structure
        doc = Document(template_path)
        style_names = [style.name for style in doc.styles]
        
        # Check that all heading styles exist
        for i in range(1, 7):
            assert f'Heading {i}' in style_names
        
        # Check basic styles
        assert 'Normal' in style_names


def test_template_heading_styles():
    """Test that template has properly configured heading styles."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        
        doc = Document(template_path)
        
        # Test heading style properties
        heading1 = doc.styles['Heading 1']
        assert heading1.font.name == 'Calibri'
        assert heading1.font.size.pt == 18
        assert heading1.font.bold is True
        
        heading2 = doc.styles['Heading 2']
        assert heading2.font.name == 'Calibri'
        assert heading2.font.size.pt == 14
        assert heading2.font.bold is True


def test_template_paragraph_styles():
    """Test that template has properly configured paragraph styles."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        
        doc = Document(template_path)
        
        # Test normal style properties
        normal_style = doc.styles['Normal']
        assert normal_style.font.name == 'Calibri'
        assert normal_style.font.size.pt == 11


def test_template_margins():
    """Test that template has correct margin settings."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        
        doc = Document(template_path)
        section = doc.sections[0]
        
        # Check 1-inch margins (914400 EMUs = 1 inch)
        assert section.top_margin.emu == 914400
        assert section.bottom_margin.emu == 914400
        assert section.left_margin.emu == 914400
        assert section.right_margin.emu == 914400


def test_template_with_converter():
    """Test using template with converter."""
    markdown_content = """# Test Heading 1

This is normal text.

## Test Heading 2

More content here.

### Test Heading 3

Final content.
"""
    
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        template_path = tmpdir_path / "template.docx"
        input_path = tmpdir_path / "test.md"
        output_path = tmpdir_path / "output.docx"
        
        # Create template and input file
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        input_path.write_text(markdown_content)
        
        # Convert using template
        converter = MarkdownToDocxConverter(reference_doc=template_path)
        result = converter.convert(input_path, output_path)
        
        assert result == output_path
        assert output_path.exists()
        
        # Verify output uses template styles
        output_doc = Document(output_path)
        headings = []
        for paragraph in output_doc.paragraphs:
            if paragraph.style and 'heading' in paragraph.style.name.lower():
                headings.append((paragraph.text, paragraph.style.name))
        
        assert len(headings) >= 3
        assert any('Test Heading 1' in h[0] for h in headings)
        assert any('Test Heading 2' in h[0] for h in headings)
        assert any('Test Heading 3' in h[0] for h in headings)


def test_template_code_style():
    """Test that template includes code block style."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        
        doc = Document(template_path)
        style_names = [style.name for style in doc.styles]
        
        # Code Block style should exist
        assert 'Code Block' in style_names
        
        code_style = doc.styles['Code Block']
        assert code_style.font.name == 'Consolas'
        assert code_style.font.size.pt == 9


def test_template_sample_content():
    """Test that template contains sample content for structure."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "template.docx"
        DocxTemplateManager.create_modern_template(template_path, add_sample=True)
        doc = Document(template_path)
        
        # Should have some sample content
        assert len(doc.paragraphs) > 0
        
        # Check for sample headings
        heading_texts = [p.text for p in doc.paragraphs if p.style and 'heading' in p.style.name.lower()]
        assert len(heading_texts) >= 3  # At least 3 sample headings


def test_template_custom_heading_styles():
    """Test template creation with custom heading styles via kwargs."""
    with TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "custom_template.docx"
        
        # Create template with custom heading font
        DocxTemplateManager.create_modern_template(
            template_path, 
            heading_font="Arial",
            heading1_size=20
        )
        
        doc = Document(template_path)
        heading1 = doc.styles['Heading 1']
        
        # Verify custom styles are applied
        assert heading1.font.name == 'Arial'
        assert heading1.font.size.pt == 20


def test_template_reusability():
    """Test that template can be used multiple times."""
    markdown_content = """# Document Title

This is test content.
"""
    
    with TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        template_path = tmpdir_path / "template.docx"
        
        # Create template once
        DocxTemplateManager.create_modern_template(template_path)
        
        # Use template multiple times
        for i in range(3):
            input_path = tmpdir_path / f"test{i}.md"
            output_path = tmpdir_path / f"output{i}.docx"
            
            input_path.write_text(markdown_content)
            
            converter = MarkdownToDocxConverter(reference_doc=template_path)
            result = converter.convert(input_path, output_path)
            
            assert result == output_path
            assert output_path.exists()


def test_template_error_handling():
    """Test template creation error handling."""
    # Test with path in a non-existent directory (should create directories)
    with TemporaryDirectory() as tmpdir:
        invalid_path = Path(tmpdir) / "nonexistent" / "template.docx"

        result = DocxTemplateManager.create_modern_template(invalid_path)
        assert result.exists()
